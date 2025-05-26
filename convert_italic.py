import re

from docx import Document

def convert_to_italic(doc_input_path: str, doc_output_path: str) -> None:
  
    doc = Document(doc_input_path)

    for para in doc.paragraphs:
        # Combine all text from runs into one string
        full_text = ''.join(run.text for run in para.runs)
        # Split into parts: normal text and bracketed parts
        parts = re.split(r'(\[.*?\])', full_text)

        # Remove all existing runs
        for run in para.runs:
            run.clear()

        # Rebuild the paragraph with new runs
        for part in parts:
            run = para.add_run()
            if part.startswith('[') and part.endswith(']'):
                run.text = part[1:-1]  # remove brackets
                run.bold = True
                run.italic = True
            else:
                run.text = part

    doc.save(doc_output_path)

if __name__ == "__main__":
    input_path = "Transcription - Crime I_2v-85.docx"
    output_path = "Transcription - Crime I_2v-85_italic.docx"

    convert_to_italic(input_path, output_path)
    
    